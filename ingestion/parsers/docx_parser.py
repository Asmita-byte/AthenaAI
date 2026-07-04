from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from backend.config import get_settings
from backend.core.logging import get_logger
from ingestion.parsers.base_parser import (
    BaseParser,
    ParsedChunk,
    ParsedDocument,
    ParsedImage,
    ParsedTable,
)

settings = get_settings()
logger = get_logger(__name__)


class DOCXParser(BaseParser):

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> ParsedDocument:
        logger.info("Parsing DOCX", filename=file_path.name)

        docx = DocxDocument(str(file_path))

        doc = ParsedDocument(
            filename=file_path.name,
            file_extension=".docx",
            total_pages=1,
            metadata=self._extract_metadata(docx),
        )

        doc.text_chunks = self._extract_text(docx, file_path)

        if self.extract_tables:
            doc.table_chunks = self._extract_tables(docx, file_path)

        if self.extract_images:
            doc.image_chunks = self._extract_images(docx, file_path)

        logger.info(
            "DOCX parsing complete",
            filename=file_path.name,
            summary=doc.summary(),
        )

        return doc

    def _extract_metadata(self, docx: DocxDocument) -> dict:
        props = docx.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created or ""),
            "modified": str(props.modified or ""),
            "total_pages": 1,
        }

    def _extract_text(self, docx: DocxDocument, file_path: Path) -> list[ParsedChunk]:
        chunks = []
        chunk_index = 0
        current_section = ""
        current_paragraphs = []

        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style.name.startswith("Heading"):
                if current_paragraphs:
                    chunks.append(
                        ParsedChunk(
                            content="\n".join(current_paragraphs),
                            chunk_type="text",
                            page_number=1,
                            chunk_index=chunk_index,
                            source_filename=file_path.name,
                            section_title=current_section,
                            metadata={"section": current_section},
                        )
                    )
                    chunk_index += 1
                    current_paragraphs = []

                current_section = text

            else:
                current_paragraphs.append(text)

        if current_paragraphs:
            chunks.append(
                ParsedChunk(
                    content="\n".join(current_paragraphs),
                    chunk_type="text",
                    page_number=1,
                    chunk_index=chunk_index,
                    source_filename=file_path.name,
                    section_title=current_section,
                    metadata={"section": current_section},
                )
            )

        return chunks

    def _extract_tables(self, docx: DocxDocument, file_path: Path) -> list[ParsedChunk]:
        chunks = []

        for table_idx, table in enumerate(docx.tables):
            rows_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_data)

            if not rows_data:
                continue

            headers = rows_data[0]
            rows = rows_data[1:]

            parsed_table = ParsedTable(
                table_index=table_idx,
                page_number=1,
                headers=headers,
                rows=rows,
                raw_text="",
            )

            chunks.append(
                ParsedChunk(
                    content=parsed_table.to_text(),
                    chunk_type="table",
                    page_number=1,
                    chunk_index=table_idx,
                    source_filename=file_path.name,
                    table=parsed_table,
                    metadata={"table_index": table_idx},
                )
            )

        return chunks

    def _extract_images(self, docx: DocxDocument, file_path: Path) -> list[ParsedChunk]:
        chunks = []
        images_dir = settings.storage_images_dir

        for img_idx, rel in enumerate(docx.part.rels.values()):
            if "image" not in rel.reltype:
                continue

            try:
                image_data = rel.target_part.blob
                image_ext = rel.target_part.content_type.split("/")[-1]

                image_filename = f"{file_path.stem}_img{img_idx}.{image_ext}"
                image_path = images_dir / image_filename

                with open(image_path, "wb") as f:
                    f.write(image_data)

                parsed_image = ParsedImage(
                    image_index=img_idx,
                    page_number=1,
                    image_path=str(image_path),
                    image_type="image",
                )

                chunks.append(
                    ParsedChunk(
                        content="Image extracted from document.",
                        chunk_type="image",
                        page_number=1,
                        chunk_index=img_idx,
                        source_filename=file_path.name,
                        image=parsed_image,
                        metadata={"image_index": img_idx},
                    )
                )

            except Exception as e:
                logger.warning(
                    "Failed to extract image",
                    img_idx=img_idx,
                    error=str(e),
                )

        return chunks
